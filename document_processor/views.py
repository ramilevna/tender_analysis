import pytesseract
from PIL import Image
from docx import Document
from django.shortcuts import render, redirect
from .forms import DocumentForm
from django.http import HttpResponse
import PyPDF2
from llama_cpp import Llama
from django.shortcuts import render


def index(request):
    return render(request, 'index.html')

def upload_file(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save()
            file_path = document.file.path

            if file_path.endswith(('png', 'jpg', 'jpeg')):
                text = pytesseract.image_to_string(Image.open(file_path), lang="rus")

            elif file_path.endswith('pdf'):
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"

            elif file_path.endswith('docx'):
                text = []
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)

                text = '\n'.join(text)

            else:
                with open(file_path, 'rb') as f:
                    text = f.read().decode('utf-8', errors='ignore')


            llm = Llama(
                model_path="Vikhr-Llama-3.2-1B-F16.gguf",
                chat_format="llama-3",
                n_ctx=4096,
                n_threads=8,
            )

            output = llm.create_chat_completion(
                messages=[
                    {"role": "system", "content": "Ты - ИИ ассистент, которому предоставляется документ, по которому необходимо написать краткое содержание предоставленного документа на русском языке. Обрати внимание, так как документы подгружаются и затем распознаются, текст в них может быть немного нарушен, соответственно, по неправильно распознанным частям документа краткое содержание делать не нужно."},
                    {
                        "role": "user",
                        "content": "Напиши краткую выжимку документа в следующем сообщении. Если найдешь в тексте контент, который подходит по этим параметрам: Наименование заказчика, Наименование проекта, Адрес(-а) расположения защищаемых объектов заказчика, Сроки выполнения проекта этапов проекта, Перечень выполняемых работ, Перечень требований по функциям проектируемой системы защиты информации, Информация о объекте(-ах) защиты, выпиши контент тематически в отдельные параграфы с указанием названия параметра. Если не найдешь, напиши просто краткое содержание."
                    },
                    {
                        "role": "user",
                        "content": text
                    }
                ],
                stream=True
            )

            generated_text = ""
            for chunk in output:
                delta = chunk['choices'][0]['delta']
                if 'role' in delta:
                    print(delta['role'], end=': ')
                elif 'content' in delta:
                    print(delta['content'], end='')
                    generated_text += delta["content"]

            summary_doc = Document()
            summary_doc.add_heading('Краткое содержание', level=1)
            summary_doc.add_paragraph(generated_text)
            docx_path = 'summary.docx'
            summary_doc.save(docx_path)

            return redirect('/download/')

    else:
        form = DocumentForm()
    return render(request, 'upload.html', {'form': form})


def download_summary(request):
    with open('summary.docx', 'rb') as docx_file:
        response = HttpResponse(docx_file.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=summary.docx'
        return response